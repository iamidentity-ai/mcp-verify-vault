#!/usr/bin/env python3
"""Render COOKBOOK.md to COOKBOOK.docx.

Cover page, auto table of contents, the architecture diagram as a figure,
code blocks styled monospace + light grey, callout boxes for any line that
starts with "Expected:".

Usage:
  python3 -m venv .venv && source .venv/bin/activate
  pip install -r scripts/requirements-render.txt
  python3 scripts/render-cookbook.py
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from markdown_it import MarkdownIt

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / 'COOKBOOK.md'
DOCX = ROOT / 'COOKBOOK.docx'
# Optional architecture diagram for the cover page. If you have one, drop it
# at docs/architecture-diagram.png and (optionally) a caption at
# docs/architecture-diagram-caption.txt. Otherwise the cover gets a
# placeholder line you can replace by hand in Word.
DIAGRAM = ROOT / 'docs' / 'architecture-diagram.png'
CAPTION_FILE = ROOT / 'docs' / 'architecture-diagram-caption.txt'
CAPTION = CAPTION_FILE.read_text().strip() if CAPTION_FILE.exists() else ''

# Document setup
doc = Document()
styles = doc.styles

# Body font
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Headings
for lvl, sz in [(1, 22), (2, 18), (3, 14), (4, 12)]:
    h = styles[f'Heading {lvl}']
    h.font.name = 'Calibri'
    h.font.size = Pt(sz)
    h.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

# Code style
code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = 'Consolas'
code_style.font.size = Pt(9)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(4)
code_style.paragraph_format.left_indent = Inches(0.25)

# Cover page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Securing Your MCP Server End-to-End')
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('IBM Verify + HashiCorp Vault')
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x42, 0x42, 0x42)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run(
    'A 30-minute hands-on walkthrough that lands a working MCP server on your '
    'laptop where every tool call is authorized by IBM Verify policy (with '
    'step-up MFA where the policy demands it) and runs under a 5-minute '
    'PostgreSQL credential that HashiCorp Vault mints fresh per call.'
)
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph().add_run('\n\n')
if DIAGRAM.exists():
    doc.add_picture(str(DIAGRAM), width=Inches(6.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(CAPTION)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
else:
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = placeholder.add_run(
        '[Architecture diagram: drop docs/architecture-diagram.png here and re-render.]'
    )
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()


def add_toc(doc):
    """Insert a Word TOC field. User must right-click and Update Field."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = (
        'Right-click and select "Update Field" to populate the table of contents.'
    )
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)


doc.add_paragraph('Table of contents', style='Heading 1')
add_toc(doc)
doc.add_page_break()


def shade(p, hex_color: str):
    """Add a background fill to a paragraph (for code blocks and callouts)."""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


md_text = MD.read_text()
md = MarkdownIt('commonmark', {'html': False, 'breaks': True})
tokens = md.parse(md_text)

i = 0
while i < len(tokens):
    t = tokens[i]
    if t.type == 'heading_open':
        level = int(t.tag[1])
        text = tokens[i + 1].content
        # Skip the document's H1 (we already have a cover page)
        if level == 1 and i < 5:
            i += 3
            continue
        doc.add_paragraph(text, style=f'Heading {min(level, 4)}')
        i += 3
    elif t.type == 'paragraph_open':
        text = tokens[i + 1].content
        if text.strip().startswith('Expected:'):
            p = doc.add_paragraph(text)
            shade(p, 'EAF4EA')
        else:
            doc.add_paragraph(text)
        i += 3
    elif t.type == 'fence':
        for line in (t.content or '').rstrip('\n').split('\n'):
            p = doc.add_paragraph(line or ' ', style='Code')
            shade(p, 'F4F4F4')
        i += 1
    elif t.type == 'bullet_list_open':
        depth = 1
        i += 1
        while depth > 0:
            tt = tokens[i]
            if tt.type == 'bullet_list_open':
                depth += 1
                i += 1
            elif tt.type == 'bullet_list_close':
                depth -= 1
                i += 1
            elif tt.type == 'list_item_open':
                j = i + 1
                while tokens[j].type != 'inline':
                    j += 1
                doc.add_paragraph(tokens[j].content, style='List Bullet')
                while tokens[i].type != 'list_item_close':
                    i += 1
                i += 1
            else:
                i += 1
    elif t.type == 'hr':
        doc.add_paragraph()
        i += 1
    else:
        i += 1

doc.save(DOCX)
print(f'Wrote {DOCX} ({DOCX.stat().st_size} bytes)')
